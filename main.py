# import os
# from docx import Document
#
# # 设置文件夹路径和输出Word文档的路径
# folder_path = '/home/zoedoet/workspace/go/doc_server'  # 替换为你的Go文件所在的文件夹路径
# output_docx_path = 'doc_server.docx'  # 输出的Word文档名
#
# # 创建一个新的Word文档
# doc = Document()
#
# # 递归遍历指定文件夹中的所有.go文件
# for root, dirs, files in os.walk(folder_path):
#     for filename in files:
#         if filename.endswith('.go'):
#             file_path = os.path.join(root, filename)
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 # 读取Go文件的内容
#                 content = file.read()
#                 # 创建一个新的段落并添加文件内容
#                 doc.add_paragraph(content)
#                 # 添加一个分隔线，以便区分不同的Go文件
#                 doc.add_paragraph('\n---\n')
#
# # 保存Word文档
# doc.save(output_docx_path)
# print(f'所有Go文件的内容已汇总到 {output_docx_path}')


import os
import subprocess
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement

# 设置文件夹路径和输出Word文档的路径
folder_path = '/home/zoedoet/workspace/go/doc_server'  # 替换为你的Go文件所在的文件夹路径
output_docx_path = 'test_summary.docx'  # 输出的Word文档名

# 创建一个新的Word文档
doc = Document()

# 在最外层文档路径下执行shell指令 tree .
tree_output = subprocess.check_output(['tree', '.', '-L', '2'], cwd=folder_path, text=True, stderr=subprocess.DEVNULL)
tree_output_lines = tree_output.strip().split('\n')

# 将tree命令的输出结果添加到Word文档的最前面
for line in tree_output_lines:
    doc.add_paragraph(line)

# 添加分隔线
doc.add_paragraph('\n---\n')

# 递归遍历指定文件夹中的所有.go文件
for root, dirs, files in os.walk(folder_path):
    for filename in files:
        if filename.endswith('.go'):
            file_path = os.path.join(root, filename)
            with open(file_path, 'r', encoding='utf-8') as file:
                # 读取Go文件的内容
                content = file.read()
                # 创建一个新的段落并添加文件内容
                doc.add_paragraph(content)

# 保存Word文档
doc.save(output_docx_path)
print(f'所有Go文件的内容已汇总到 {output_docx_path}')
